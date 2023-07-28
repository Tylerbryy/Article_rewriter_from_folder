import subprocess
import pkg_resources

REQUIRED_PACKAGES = [
    'python-dotenv',
    'openai',
    'python-docx',
    'tqdm',
    'ttkthemes'
]

for package in REQUIRED_PACKAGES:
    try:
        dist = pkg_resources.get_distribution(package)
    except pkg_resources.DistributionNotFound:
        subprocess.call(['pip', 'install', package])


import os
import openai
from docx import Document
from dotenv import load_dotenv
from time import sleep
from openai import OpenAIError
from tkinter import filedialog
from tkinter import messagebox
from tkinter import Tk, Label, Entry, StringVar
from tkinter.ttk import Progressbar, Style, Button
from ttkthemes import ThemedStyle
from threading import Thread
import sv_ttk


# Load OpenAI key
load_dotenv()
openai.api_key = os.getenv("open_api_key")

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return ' '.join(full_text)

def write_docx(file_path, content):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(file_path)

def generate_rewritten_content_with_retry(content: str):
    retries = 0
    while retries < 10:  # Maximum 10 retries
        try:
            messages = []
            messages.append({"role": "user", "content": f"""Please rewrite the following article in a different way and make it longer : "{content}" """})
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=messages
            )
            return response.choices[0].message.content
        except OpenAIError as e:
            print("API Error:", e)
            sleep(1)
            print("Retrying in 15 seconds...")
            sleep(15)
            retries += 1

class GUI:
    def __init__(self, master):
        self.master = master
        master.title("OpenAI Article Rewriter")

        sv_ttk.set_theme("dark")  # Use the 'arc' theme

        self.label_input = Label(master, text="Input Directory")
        self.label_input.pack()

        self.input_dir = StringVar()
        self.entry_input = Entry(master, textvariable=self.input_dir)
        self.entry_input.pack()

        self.button_browse_input = Button(master, text="Browse", command=self.browse_input)
        self.button_browse_input.pack()

        self.label_output = Label(master, text="Output Directory")
        self.label_output.pack()

        self.output_dir = StringVar()
        self.entry_output = Entry(master, textvariable=self.output_dir)
        self.entry_output.pack()

        self.button_browse_output = Button(master, text="Browse", command=self.browse_output)
        self.button_browse_output.pack()

        # Progress bar
        self.progress_bar = Progressbar(master, length=500, mode='determinate')
        self.progress_bar.pack(pady=10)

        self.button_start = Button(master, text="Start", command=self.start_rewrite_articles_thread)
        self.button_start.pack()

    def browse_input(self):
        self.input_dir.set(filedialog.askdirectory())

    def browse_output(self):
        self.output_dir.set(filedialog.askdirectory())

    def start_rewrite_articles_thread(self):
        # Disable start button
        self.button_start['state'] = 'disabled'
        self.progress_bar['value'] = 0

        # Create and start new thread running the rewrite_articles() function
        thread = Thread(target=self.rewrite_articles)
        thread.start()

    def rewrite_articles(self):
        input_dir = self.input_dir.get()
        output_dir = self.output_dir.get()

        try:
            if not os.path.isdir(input_dir):
                raise Exception(f"Input directory {input_dir} does not exist.")
            if not os.path.isdir(output_dir):
                os.makedirs(output_dir)

            file_count = len(os.listdir(input_dir))
            self.progress_bar['maximum'] = file_count

            for i, file_name in enumerate(os.listdir(input_dir), start=1):
                # Only process docx files
                if file_name.endswith('.docx'):
                    # Read original content
                    original_content = read_docx(os.path.join(input_dir, file_name))

                    # Generate rewritten content
                    rewritten_content = generate_rewritten_content_with_retry(original_content)

                    # Write rewritten content to new file in output directory
                    write_docx(os.path.join(output_dir, file_name), rewritten_content)

                # Update progress bar
                self.progress_bar['value'] = i
                self.master.update()

            messagebox.showinfo("Info", "Rewriting complete")

        except Exception as e:
            messagebox.showerror("Error", str(e))

        finally:
            # Enable start button
            self.button_start['state'] = 'normal'

root = Tk()
my_gui = GUI(root)
root.mainloop()
